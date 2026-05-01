from __future__ import annotations

import json
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

import docframe as df
from docframe.cli import collect_paths
from docx import Document as DocxDocument
from openpyxl import Workbook
from PIL import Image
from pypdf import PdfWriter


class DocFrameTests(unittest.IsolatedAsyncioTestCase):
    async def test_processes_csv_into_table_chunk(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "users.csv"
            path.write_text("name,email\nAda,ada@example.internal\n\n", encoding="utf-8")

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "csv")
        self.assertEqual(result.metadata.row_count, 1)
        self.assertEqual(result.chunks[0].type, "table")
        self.assertEqual(result.chunks[0].rows[0]["name"], "Ada")

    async def test_processes_docx_text(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "brief.docx"
            document = DocxDocument()
            document.add_paragraph("DocFrame reads Word documents.")
            document.save(path)

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "docx")
        self.assertIn("Word documents", result.text)

    async def test_processes_legacy_doc_as_metadata_only(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "legacy.doc"
            path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "doc")
        self.assertEqual(result.chunks[0].type, "metadata")
        self.assertTrue(result.warnings)

    async def test_processes_ooxml_file_with_doc_extension(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "misnamed.doc"
            document = DocxDocument()
            document.add_paragraph("DocFrame recovers misnamed OOXML Word files.")
            document.save(path)

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "doc")
        self.assertIn("misnamed OOXML", result.text)
        self.assertTrue(result.warnings)

    async def test_processes_xlsx_into_sheet_tables(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "book.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Users"
            sheet.append(["name", "score"])
            sheet.append(["Grace", 10])
            workbook.save(path)

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "xlsx")
        self.assertEqual(result.metadata.sheet_count, 1)
        self.assertEqual(result.chunks[0].sheet, "Users")
        self.assertEqual(result.chunks[0].rows[0]["name"], "Grace")

    async def test_processes_pdf_metadata(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "blank.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            with path.open("wb") as handle:
                writer.write(handle)

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "pdf")
        self.assertEqual(result.metadata.page_count, 1)

    async def test_processes_png_image_metadata(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "image.png"
            Image.new("RGB", (16, 9), color="white").save(path)

            result = await df.DocFrame().process(path)

        self.assertEqual(result.metadata.document_type, "image")
        self.assertEqual(result.metadata.width, 16)
        self.assertEqual(result.metadata.height, 9)
        self.assertEqual(result.chunks[0].type, "image")

    async def test_pipeline_steps_can_enrich_results(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "users.csv"
            path.write_text("name\nAda\n", encoding="utf-8")

            async def mark(result: df.DocumentResult) -> df.DocumentResult:
                result.metadata.extra["pipeline"] = "ran"
                return result

            framework = df.DocFrame()
            framework.use(mark)
            result = await framework.process(path)

        self.assertEqual(result.metadata.extra["pipeline"], "ran")

    async def test_directory_collection_filters_supported_formats(self) -> None:
        with TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "a.csv").write_text("x\n1\n", encoding="utf-8")
            (root / "ignore.txt").write_text("nope", encoding="utf-8")

            paths = collect_paths(root, recursive=False, framework=df.DocFrame())

        self.assertEqual([path.name for path in paths], ["a.csv"])

    async def test_result_serializes_to_json(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "users.csv"
            path.write_text("name\nAda\n", encoding="utf-8")

            result = await df.DocFrame().process(path)
            payload = json.loads(result.model_dump_json())

        self.assertEqual(payload["metadata"]["document_type"], "csv")

    async def test_process_many_can_return_structured_errors(self) -> None:
        with TemporaryDirectory() as directory:
            path = Path(directory) / "broken.pdf"
            path.write_bytes(b"not really a pdf")

            results = await df.DocFrame().process_many([path], continue_on_error=True)

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0].metadata.document_type, "pdf")
        self.assertEqual(results[0].chunks, [])
        self.assertTrue(results[0].errors)

    async def test_process_many_preserves_input_order_with_bounded_concurrency(self) -> None:
        with TemporaryDirectory() as directory:
            root = Path(directory)
            first = root / "first.csv"
            second = root / "second.csv"
            first.write_text("name\nAda\n", encoding="utf-8")
            second.write_text("name\nGrace\n", encoding="utf-8")

            framework = df.DocFrame(options=df.ProcessingOptions(max_concurrency=1))
            results = await framework.process_many([first, second])

        self.assertEqual([result.metadata.filename for result in results], ["first.csv", "second.csv"])


if __name__ == "__main__":
    unittest.main()
